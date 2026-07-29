import streamlit as st
from docx import Document
import requests
from bs4 import BeautifulSoup
import json
import io
import re
import time
import base64
import os
import glob
from datetime import datetime

HISTORY_DIR = "history"

def get_base64_image(image_path):
    try:
        with open(image_path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    except Exception:
        return ""

bg_base64 = get_base64_image("assets/bg.jpg")

# ==================== CONFIGURATION ====================

EXECUTOR_MODEL = "qwen3.5:4b"       # Local — fast extraction & summarization
EVALUATOR_MODEL = "gpt-oss:20b-cloud"  # Cloud — quality final rewrite

st.set_page_config(page_title="Resume optimization Agent", layout="wide")

st.markdown(f"""
<style>
    /* Modern Background & Font */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;800&display=swap');
    
    html, body, [class*="css"]  {{
        font-family: 'Inter', sans-serif;
    }}
    
    .stApp {{
        background-color: #f8fafc;
        background-image: url("data:image/jpeg;base64,{bg_base64}");
        background-size: cover;
        background-position: center;
        background-attachment: fixed;
    }}
    
    /* Sleek gradient buttons */
    .stButton>button {{
        background: linear-gradient(135deg, #3b82f6 0%, #8b5cf6 100%);
        color: white !important;
        font-weight: 600;
        border-radius: 12px;
        border: none;
        box-shadow: 0 4px 14px 0 rgba(139, 92, 246, 0.39);
        transition: all 0.3s ease;
        padding: 0.5rem 1rem;
    }}
    .stButton>button:hover {{
        transform: translateY(-2px);
        box-shadow: 0 6px 20px 0 rgba(139, 92, 246, 0.5);
    }}
    
    /* Clean inputs */
    .stTextInput>div>div>input, .stTextArea>div>div>textarea {{
        border-radius: 10px;
        border: 1px solid rgba(226, 232, 240, 0.5);
        box-shadow: inset 0 2px 4px 0 rgba(0, 0, 0, 0.02);
        background: rgba(255, 255, 255, 0.4) !important;
        backdrop-filter: blur(10px);
    }}
    .stTextInput>div>div>input:focus, .stTextArea>div>div>textarea:focus {{
        border-color: #8b5cf6;
        box-shadow: 0 0 0 1px #8b5cf6;
        background: rgba(255, 255, 255, 0.2) !important;
    }}

    /* Cards / Expanders */
    .streamlit-expanderHeader {{
        background-color: rgba(255, 255, 255, 0.4);
        backdrop-filter: blur(10px);
        border-radius: 12px;
        border: 1px solid rgba(226, 232, 240, 0.5);
        font-weight: 600;
        color: #1e293b;
    }}
    
    /* Sidebar specific */
    [data-testid="stSidebar"] {{
        background-color: rgba(255, 255, 255, 0.4) !important;
        backdrop-filter: blur(12px);
    }}
    
    /* Headers */
    h1, h2, h3 {{
        color: #0f172a;
        font-weight: 800;
        text-shadow: 0 1px 2px rgba(255,255,255,0.8);
    }}
</style>
""", unsafe_allow_html=True)

# ==================== OLLAMA HELPERS ====================

def check_ollama_status(base_url: str) -> dict:
    """Check if Ollama is running and which required models are available."""
    status = {"online": False, "models": {}}
    try:
        resp = requests.get(f"{base_url}/api/tags", timeout=5)
        if resp.status_code == 200:
            status["online"] = True
            available = [m["name"] for m in resp.json().get("models", [])]
            for model in [EXECUTOR_MODEL, EVALUATOR_MODEL]:
                # Match by base name (e.g. "qwen3.5:4b" matches "qwen3.5:4b")
                status["models"][model] = any(
                    model == a or model == a.split(":")[0] for a in available
                )
    except requests.ConnectionError:
        pass
    return status


def ollama_chat(
    model: str,
    messages: list,
    base_url: str,
    temperature: float = 0.3,
    max_tokens: int = 4096,
    think: bool | None = None,
) -> str:
    """
    Call Ollama's local REST API and return the response text.
    
    Args:
        think: Explicitly control thinking mode. Set False to disable
               Qwen3.5's internal reasoning loop which can cause hangs.
    """
    payload = {
        "model": model,
        "messages": messages,
        "stream": False,
        "options": {
            "temperature": temperature,
            "num_predict": max_tokens,
        },
    }
    # Ollama supports a top-level "think" flag for Qwen3-family models
    if think is not None:
        payload["think"] = think

    response = requests.post(
        f"{base_url}/api/chat",
        json=payload,
        timeout=180,
    )
    response.raise_for_status()
    msg = response.json()["message"]
    content = msg.get("content", "")

    # Fallback: some models intermittently put output in the thinking
    # field instead of content — extract JSON from there if needed
    if not content.strip() and msg.get("thinking"):
        content = msg["thinking"]

    return content


def strip_think_tags(text: str) -> str:
    """Remove Qwen3.5 <think>...</think> reasoning blocks from output."""
    return re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL).strip()


def extract_json_from_text(text: str) -> dict | None:
    """
    Robustly extract a JSON object from LLM output.
    Handles: think tags, markdown fences, surrounding commentary.
    """
    # Strip any leaked <think> blocks first
    text = strip_think_tags(text)

    # Try direct parse first
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Try to find JSON block in markdown fences
    fence_match = re.search(r"```(?:json)?\s*\n?(.*?)\n?```", text, re.DOTALL)
    if fence_match:
        try:
            return json.loads(fence_match.group(1).strip())
        except json.JSONDecodeError:
            pass

    # Try to find the outermost { ... } block
    brace_match = re.search(r"\{.*\}", text, re.DOTALL)
    if brace_match:
        try:
            return json.loads(brace_match.group(0))
        except json.JSONDecodeError:
            pass

    return None


# ==================== DOCUMENT HELPERS ====================

def extract_text_from_jd_url(url: str, base_url: str) -> str:
    """Scrapes raw text from a Job Description URL, with aggressive noise removal."""
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }
        response = requests.get(url, headers=headers, timeout=10)
        if response.status_code != 200:
            return f"Error: Unable to fetch page (Status Code: {response.status_code})"

        soup = BeautifulSoup(response.text, "html.parser")

        # Remove non-content elements aggressively
        for tag in soup(["script", "style", "nav", "header", "footer",
                         "aside", "form", "button", "noscript", "svg",
                         "iframe", "img", "input", "select", "textarea"]):
            tag.decompose()

        # Try to find the main content container first
        main_content = (
            soup.find("div", class_=re.compile(r"description|job.?body|job.?detail", re.I))
            or soup.find("article")
            or soup.find("main")
            or soup
        )

        # Extract text and clean spacing
        text = main_content.get_text(separator="\n")
        lines = [line.strip() for line in text.splitlines()]
        # Filter out short junk lines (nav items, button labels, etc.)
        meaningful_lines = [
            line for line in lines
            if len(line) > 30 or any(
                kw in line.lower()
                for kw in ["requirement", "responsibilit", "qualif", "experience",
                           "skill", "degree", "bachelor", "master", "year"]
            )
        ]
        raw_text = "\n".join(meaningful_lines)

        if not raw_text.strip():
            # Fallback: use all text if filtering removed everything
            raw_text = "\n".join(line for line in lines if line)

        # Use Qwen to extract just the JD from noisy scraped text
        return _clean_jd_with_llm(raw_text[:8000], base_url)
    except Exception as e:
        return f"Error parsing URL: {str(e)}"


def _clean_jd_with_llm(raw_text: str, base_url: str) -> str:
    """Use Qwen3.5:4b to extract the actual job description from noisy scraped text."""
    try:
        cleaned = ollama_chat(
            model=EXECUTOR_MODEL,
            messages=[
                {"role": "system", "content": "You extract job descriptions from noisy web page text. Return ONLY the cleaned job description content — job title, company, responsibilities, requirements, qualifications, and benefits. Remove all navigation text, login prompts, ads, and UI elements. Do not add commentary."},
                {"role": "user", "content": f"Extract the job description from this scraped web page:\n\n{raw_text}\n\n/no_think"},
            ],
            base_url=base_url,
            temperature=0.1,
            max_tokens=3072,
            think=False,
        )
        return strip_think_tags(cleaned) if cleaned.strip() else raw_text
    except Exception:
        # If LLM cleanup fails, return the raw text as fallback
        return raw_text


def _is_rewritable_paragraph(paragraph) -> bool:
    """Detect if a paragraph is a meaningful experience bullet point that should be rewritten."""
    text = paragraph.text.strip()
    if not text or len(text) < 30:
        return False
        
    # Exclude headers, titles, dates, contact info, and known short lines
    if re.search(r'\d{4}\s*[-–]\s*\d{4}', text):
        return False
    if "jinxl" in text or "Lyken" in text or "github" in text.lower():
        return False
        
    # Identify project subtitle / technical skills lines that rely on tabs
    if '\t' in paragraph.text and ',' in text:
        return True
        
    # If it has a manual bullet, it's definitely a bullet point
    if text.startswith(('•', '-', '', '·', '*')):
        return True
        
    # If it uses a built-in list style
    style_name = (paragraph.style.name or "").lower()
    if 'list' in style_name:
        return True
        
    # If it's a long paragraph without dates, it might be an unbulleted experience summary
    if len(text) > 100:
        # Check if it looks like a title (contains a date range like "2024 - 2025")
        if re.search(r'\b20\d{2}\s*[-–—]\s*(20\d{2}|Present)\b', text, re.IGNORECASE):
            return False
        return True
        
    return False


def extract_paragraphs_from_docx(docx_file) -> list:
    """
    Extracts content paragraphs from the uploaded resume file.
    Strictly extracts ONLY actual experience descriptions (bullet points).
    Protects headers, titles, dates, contact info, and metadata by skipping them.
    """
    doc = Document(docx_file)
    paragraphs = []
    for p in doc.paragraphs:
        if _is_rewritable_paragraph(p):
            paragraphs.append(p.text.strip())
    return paragraphs


def apply_revisions_to_docx(original_docx_file, mapping_dict: dict) -> io.BytesIO:
    """
    Applies revisions to a copy of the original docx while fully preserving
    paragraph-level formatting (style, alignment, indentation, numbering) and
    character-level formatting (font, size, color, bold, italic, etc.).

    Strategy: for each matched paragraph, remove all existing runs from the
    XML tree, then insert a single new run carrying a deep-copied rPr
    (run properties) element from the original first run.
    """
    from copy import deepcopy
    from docx.oxml.ns import qn

    doc = Document(original_docx_file)

    # Build lookup from original text -> revised text
    lookup = {
        item["original"].strip(): item["revised"].strip()
        for item in mapping_dict.get("mappings", [])
    }

    for paragraph in doc.paragraphs:
        cleaned_para_text = paragraph.text.strip()
        if cleaned_para_text not in lookup:
            continue

        new_text = lookup[cleaned_para_text]
        p_elem = paragraph._element

        # Capture the run-properties XML from the first run (if any)
        old_runs = p_elem.findall(qn("w:r"))
        rPr_copy = None
        if old_runs:
            rPr = old_runs[0].find(qn("w:rPr"))
            if rPr is not None:
                rPr_copy = deepcopy(rPr)

        # Remove all existing runs from the paragraph XML
        for r in old_runs:
            p_elem.remove(r)

        # Create a new run with the revised text
        new_run = paragraph.add_run(new_text)

        # Apply the deep-copied formatting from the original first run
        if rPr_copy is not None:
            existing_rPr = new_run._element.find(qn("w:rPr"))
            if existing_rPr is not None:
                new_run._element.remove(existing_rPr)
            new_run._element.insert(0, rPr_copy)

    # Save to a dynamic buffer
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# ==================== STAGE 1: EXECUTOR (Qwen3.5:4b) ====================

def stage1_extract_intelligence(resume_paragraphs: list, job_description: str, base_url: str) -> dict | None:
    """
    Stage 1 — Local Executor (Qwen3.5:4b)
    Extracts structured intelligence from the resume and JD:
    - Summarized resume skills
    - Required JD skills and keywords
    - Gaps between resume and JD
    """
    system_prompt = """You are a resume analysis assistant. Extract and summarize key information from a resume and a job description. Return ONLY a valid JSON object."""

    user_prompt = f"""Extract skills and keywords from the resume and job description below.

Return a JSON object with these exact keys:
- "resume_skills": list of all skills, tools, frameworks from the resume
- "resume_experience_summary": 2-3 sentence summary of the candidate
- "jd_required_skills": required skills from the JD
- "jd_preferred_skills": preferred/nice-to-have skills from the JD
- "jd_keywords": important ATS keywords and phrases from the JD
- "gaps": skills in the JD that are missing from the resume

### RESUME PARAGRAPHS:
{json.dumps(resume_paragraphs, indent=2)}

### JOB DESCRIPTION:
{job_description}

Respond with ONLY the JSON object, no other text. /no_think"""

    max_retries = 2
    for attempt in range(max_retries + 1):
        try:
            raw_response = ollama_chat(
                model=EXECUTOR_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                base_url=base_url,
                temperature=0.1,
                max_tokens=3072,
                think=False,
            )

            result = extract_json_from_text(raw_response)
            if result and "resume_skills" in result and "jd_keywords" in result:
                return result

            if attempt < max_retries:
                st.warning(f"Stage 1: Retrying extraction (attempt {attempt + 2}/{max_retries + 1})...")
        except Exception as e:
            if attempt < max_retries:
                st.warning(f"Stage 1: Retrying after error (attempt {attempt + 2}/{max_retries + 1})...")
            else:
                st.error(f"Stage 1 failed: {e}")

    return None


# ==================== STAGE 2: EVALUATOR (gpt-oss:20b-cloud) ====================

def stage2_rewrite_resume(resume_paragraphs: list, intelligence: dict, base_url: str) -> dict | None:
    """
    Stage 2 — Cloud Evaluator (gpt-oss:20b-cloud)
    Uses the intelligence from Stage 1 to produce high-quality ATS-optimized rewrites.
    Processes paragraphs in batches to avoid output truncation on long resumes.
    """
    system_prompt = """You are an expert ATS (Applicant Tracking System) optimization agent. You will receive:
1. Original resume paragraphs (bullet points or skill lines)
2. An intelligence package containing: the candidate's skills, the JD's required skills & keywords, and identified gaps

### GOAL:
1. MAXIMIZE ATS KEYWORD MATCHING for the JD provided.
2. ELIMINATE REASON FOR REJECTION due to keyword stuffing, fluff, or unnatural tail-ended phrase lists.
3. MAINTAIN HIGH IMPACT using the standard formula: [Action Verb] + [Technical Method / Data Scale] + [Business Outcome / Strategic Impact].
4. CONCISENESS (ONE-PAGE LIMIT): To keep the resume to exactly one page, your revised text MUST NOT be significantly longer than the original text. Be extremely punchy and concise.

### TARGET JOB DESCRIPTION KEYWORDS TO INTEGRATE NATURALLY:
(These will be provided in the INTELLIGENCE PACKAGE below. The Technical skills are the core area for keyword matching.)

### STRICT RULES:
1. TRUTHFULNESS: Do NOT invent false information, change degree titles, or alter numerical metrics (e.g., keep "Kaggle Rank #1", "47k+ observations"). Only rephrase and emphasize what the candidate already has.
2. SHOW, DON'T TELL: Show behavioral traits (e.g., leadership, resilience, learner mindset) through concrete actions and project outcomes rather than explicitly stating the buzzwords.
3. SKILL SUBLINES (STRICT FORMATTING): If a paragraph contains a tab character (\t) or vertical bar (|), you MUST preserve the EXACT structure and formatting, including the \t character. Your ONLY job for these lines is to reorder or substitute the HARD TECHNICAL SKILLS to heavily emphasize the JD keywords.
- CRITICAL: Do NOT modify non-technical elements (like locations e.g. "Vancouver, Canada", or project names).
- CRITICAL: ONLY inject HARD technical skills (tools, languages, algorithms). Do NOT inject soft skills, cultural buzzwords, or business jargon (e.g. "diversity", "coaching", "leadership"). Do NOT turn it into a sentence.
4. GAP BRIDGING: For identified gaps where the candidate has *adjacent* experience, rephrase to highlight the connection.
5. ONLY include paragraphs that actually need revision. If a paragraph is already well-optimized, omit it.

You MUST return ONLY a valid JSON object:
{
  "mappings": [
    {
      "original": "Exact original text...",
      "revised": "ATS-optimized rewrite..."
    }
  ]
}

Return ONLY the JSON. No explanations, no markdown fences, no conversational text."""

    intelligence_block = f"""**Candidate's Skills:** {json.dumps(intelligence.get("resume_skills", []))}
**Experience Summary:** {intelligence.get("resume_experience_summary", "N/A")}
**JD Required Skills:** {json.dumps(intelligence.get("jd_required_skills", []))}
**JD Preferred Skills:** {json.dumps(intelligence.get("jd_preferred_skills", []))}
**ATS Keywords:** {json.dumps(intelligence.get("jd_keywords", []))}
**Identified Gaps:** {json.dumps(intelligence.get("gaps", []))}"""

    # Small batches — gpt-oss cloud backend has a fixed output token limit
    BATCH_SIZE = 2
    all_mappings = []
    batches = [
        resume_paragraphs[i : i + BATCH_SIZE]
        for i in range(0, len(resume_paragraphs), BATCH_SIZE)
    ]

    for batch_idx, batch in enumerate(batches):
        batch_label = f"Batch {batch_idx + 1}/{len(batches)}"
        st.write(f"  Processing {batch_label} ({len(batch)} paragraphs)...")

        user_prompt = f"""### INTELLIGENCE PACKAGE:
{intelligence_block}

### ORIGINAL RESUME PARAGRAPHS TO OPTIMIZE:
{json.dumps(batch, indent=2)}

Return ONLY the JSON object with the "mappings" array. /no_think"""

        batch_result = _call_evaluator_with_retries(
            system_prompt, user_prompt, batch_label, base_url
        )
        if batch_result is None:
            return None  # Abort on any batch failure

        all_mappings.extend(batch_result.get("mappings", []))

    return {"mappings": all_mappings}


def _call_evaluator_with_retries(
    system_prompt: str, user_prompt: str, label: str, base_url: str
) -> dict | None:
    """Call the evaluator model with retries and debug logging."""
    max_retries = 2
    for attempt in range(max_retries + 1):
        try:
            raw_response = ollama_chat(
                model=EVALUATOR_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                base_url=base_url,
                temperature=0.2,
                max_tokens=8192,
                think=False,
            )

            result = extract_json_from_text(raw_response)
            if result and "mappings" in result:
                return result

            # Debug: detect truncation vs. other parse failure
            stripped = raw_response.rstrip()
            is_truncated = stripped and stripped[-1] not in ("}", "]")
            diag = "TRUNCATED output" if is_truncated else "malformed JSON"
            preview = raw_response[-300:] if raw_response else "(empty)"

            if attempt < max_retries:
                st.warning(
                    f"Stage 2 ({label}): {diag}, retrying "
                    f"(attempt {attempt + 2}/{max_retries + 1})..."
                )
                time.sleep(2)  # Brief pause before retry to reduce contention
            else:
                st.error(
                    f"Stage 2 ({label}): Could not parse response — {diag}.\n\n"
                    f"**Response tail:** `{preview}`"
                )
        except requests.exceptions.ReadTimeout:
            if attempt < max_retries:
                st.warning(
                    f"Stage 2 ({label}): Timed out, retrying "
                    f"(attempt {attempt + 2}/{max_retries + 1})..."
                )
            else:
                st.error(f"Stage 2 ({label}): Request timed out after all retries.")
        except Exception as e:
            if attempt < max_retries:
                st.warning(
                    f"Stage 2 ({label}): Error, retrying "
                    f"(attempt {attempt + 2}/{max_retries + 1})..."
                )
            else:
                st.error(f"Stage 2 ({label}) failed: {e}")

    return None


# ==================== STREAMLIT UI ====================

st.markdown("""
<div style="text-align: center; padding-bottom: 2rem; padding-top: 1rem;">
    <h1 style="font-size: 3rem; margin-bottom: 0.5rem;">Resume optimization Agent</h1>
    <p style="color: #475569; font-size: 1.1rem; font-weight: 500;">Two-stage AI pipeline · Qwen extracts intelligence · GPT ensures quality · Fully via Ollama</p>
</div>
""", unsafe_allow_html=True)

# ---- Sidebar ----
with st.sidebar:
    st.header("System Status")
    
    ollama_url = st.text_input("Ollama Server URL", value="http://localhost:11434", help="If deploying to Streamlit Cloud, use Ngrok to expose your local Ollama port (e.g. https://xxxx.ngrok-free.app)")

    ollama_status = check_ollama_status(ollama_url)

    if ollama_status["online"]:
        st.success("🟢 Ollama is running")
    else:
        st.error("🔴 Ollama is offline — run `ollama serve`")

    st.markdown("---")
    st.markdown("---")
    st.markdown("### Past Optimizations")
    
    history_files = sorted(glob.glob(os.path.join(HISTORY_DIR, "*_meta.json")), reverse=True)
    if not history_files:
        st.info("No past optimizations found.")
    else:
        for mf in history_files:
            try:
                with open(mf, "r") as f:
                    meta = json.load(f)
                ts = meta.get("timestamp")
                if not ts:
                    continue
                dt = datetime.strptime(ts, "%Y%m%d_%H%M%S")
                label = dt.strftime("%b %d, %Y - %I:%M %p")
                
                if st.button(f"📄 {label}", key=f"btn_{ts}", use_container_width=True):
                    st.session_state["intelligence"] = meta["intelligence"]
                    st.session_state["optimization_mappings"] = meta["mappings"]
                    st.session_state["history_docx_path"] = os.path.join(HISTORY_DIR, f"{ts}_original.docx")
            except Exception:
                pass

# ---- Core Inputs ----
col1, col2 = st.columns(2)

with col1:
    st.header("Upload Original Resume")
    uploaded_file = st.file_uploader("Upload your resume (.docx only)", type=["docx"])

with col2:
    st.header("Job Description Source")
    jd_source = st.radio("Provide Job Description via:", ["Web URL", "Paste Text"])

    job_description_text = ""
    if jd_source == "Web URL":
        jd_url = st.text_input("Enter Job Description Link (e.g., LinkedIn, Greenhouse, Indeed)")
        if jd_url:
            with st.spinner("Scraping Job Description..."):
                job_description_text = extract_text_from_jd_url(jd_url, ollama_url)
                if "Error" in job_description_text:
                    st.error(job_description_text)
                    job_description_text = ""
                else:
                    st.success("Successfully scraped Job Description!")
                    with st.expander("Show Scraped Text"):
                        st.write(job_description_text[:1000] + "...")
    else:
        job_description_text = st.text_area("Paste Job Description Text here...", height=200)

# ---- Optimization Pipeline ----
if st.button("🚀 Optimize My Resume for ATS Screening", use_container_width=True):
    if not ollama_status["online"]:
        st.error("Ollama is not running. Please start it with `ollama serve` and refresh.")
    elif not all(ollama_status["models"].values()):
        missing = [m for m, v in ollama_status["models"].items() if not v]
        st.error(f"Missing model(s): {', '.join(missing)}. Pull them with `ollama pull <model>`.")
    elif not uploaded_file:
        st.error("Please upload a .docx resume file first.")
    elif not job_description_text:
        st.error("Please provide a Job Description.")
    else:
        # Step 1: Extract paragraphs
        paragraphs = extract_paragraphs_from_docx(uploaded_file)

        # Step 2: Stage 1 — Executor
        with st.status("Stage 1: Analyzing resume & JD with Qwen...", expanded=True) as status1:
            st.write(f"Sending {len(paragraphs)} resume paragraphs + JD to `Qwen`...")
            intelligence = stage1_extract_intelligence(paragraphs, job_description_text, ollama_url)

            if intelligence:
                status1.update(label="Stage 1 Complete: Intelligence extracted", state="complete")
                st.session_state["intelligence"] = intelligence
            else:
                status1.update(label="Stage 1 Failed", state="error")
                st.error("Could not extract intelligence from the resume and JD. Please try again.")
                st.stop()

        # Show intelligence preview
        with st.expander("Intelligence Preview — What the AI Sees", expanded=True):
            intel_col1, intel_col2, intel_col3 = st.columns(3)
            with intel_col1:
                st.markdown("**Your Skills**")
                for skill in intelligence.get("resume_skills", []):
                    st.markdown(f"- {skill}")
            with intel_col2:
                st.markdown("**JD Keywords & Required Skills**")
                for kw in intelligence.get("jd_keywords", []):
                    st.markdown(f"- 🔑 {kw}")
                for skill in intelligence.get("jd_required_skills", []):
                    st.markdown(f"- ⚡ {skill}")
            with intel_col3:
                st.markdown("**Identified Gaps**")
                for gap in intelligence.get("gaps", []):
                    st.markdown(f"- ⚠️ {gap}")

            if intelligence.get("resume_experience_summary"):
                st.info(f"**Experience Summary:** {intelligence['resume_experience_summary']}")

        # Step 3: Stage 2 — Evaluator
        with st.status("Stage 2: Rewriting with GPT evaluator...", expanded=True) as status2:
            st.write(f"Sending intelligence + paragraphs to `GPT` for quality rewrite...")
            optimizer_response = stage2_rewrite_resume(paragraphs, intelligence, ollama_url)

            if optimizer_response and "mappings" in optimizer_response:
                status2.update(label="Stage 2 Complete: Resume optimized", state="complete")
                st.session_state["optimization_mappings"] = optimizer_response
                st.session_state.pop("history_docx_path", None)
                st.success("🎉 Optimization complete!")
                
                # Save to history
                try:
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    uploaded_file.seek(0)
                    original_path = os.path.join(HISTORY_DIR, f"{timestamp}_original.docx")
                    with open(original_path, "wb") as f:
                        f.write(uploaded_file.read())
                    
                    meta_path = os.path.join(HISTORY_DIR, f"{timestamp}_meta.json")
                    with open(meta_path, "w") as f:
                        json.dump({
                            "timestamp": timestamp,
                            "intelligence": intelligence,
                            "mappings": optimizer_response
                        }, f)
                except Exception as e:
                    st.warning(f"Could not save history: {e}")
            else:
                status2.update(label="Stage 2 Failed", state="error")
                st.error("Could not produce optimized rewrites. Please try again.")

# ---- Output Display & Download ----
if "optimization_mappings" in st.session_state:
    st.header("Compare & Download Revisions")

    mappings = st.session_state["optimization_mappings"]["mappings"]

    if not mappings:
        st.info("No revisions were suggested — your resume is already well-optimized for this JD!")
    else:
        # Display before / after changes
        for index, item in enumerate(mappings):
            with st.expander(f"Revision #{index + 1}: {item['original'][:60]}..."):
                c_orig, c_rev = st.columns(2)
                with c_orig:
                    st.markdown("**Original:**")
                    st.write(item["original"])
                with c_rev:
                    st.markdown("**ATS Optimized:**")
                    st.success(item["revised"])

        # Recompile DOCX for download
        if "history_docx_path" in st.session_state and os.path.exists(st.session_state["history_docx_path"]):
            with open(st.session_state["history_docx_path"], "rb") as f:
                source_docx = io.BytesIO(f.read())
        else:
            uploaded_file.seek(0)
            source_docx = uploaded_file
            
        revised_docx_buffer = apply_revisions_to_docx(
            source_docx, st.session_state["optimization_mappings"]
        )

        st.download_button(
            label="📥 Download Tailored Resume (.docx)",
            data=revised_docx_buffer,
            file_name="ATS_Optimized_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )